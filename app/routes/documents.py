import io
from flask import Blueprint, render_template, request, redirect, url_for, abort, flash
from app.extensions import db
from app.models.project import Project
from app.models.document import Document
from app.services.annotation_service import annotate_document
from app.services.highlighter import build_highlighted_text
from app.services.export_service import export_csv
from app.services.export_service import export_excel

import docx

documents_bp = Blueprint("documents", __name__, url_prefix="/documents")


def extract_text_from_docx(file_stream):
    doc = docx.Document(file_stream)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text)


@documents_bp.route("/upload/<int:project_id>", methods=["GET", "POST"])
def upload_document(project_id):
    project = Project.query.get(project_id)
    if not project:
        abort(404)

    if request.method == "POST":
        name = request.form.get("name")
        content = request.form.get("content", "").strip()

        file = request.files.get("file")

        #проверяем на прикрепление файла
        if file and file.filename != '':
            filename = file.filename.lower()

            #проверка на тип файла
            if not (filename.endswith('.docx') or filename.endswith('.txt')):
                flash("Неподдерживаемый формат файла! Допускаются только документы .docx и .txt", "error")
                return render_template("documents/upload.html", project=project)

            try:
                # Считываем бинарный поток в оперативную память
                file_stream = io.BytesIO(file.read())

                if filename.endswith('.docx'):
                    content = extract_text_from_docx(file_stream)
                elif filename.endswith('.txt'):
                    # Читаем сырые байты .txt файла и пробуем декодировать в UTF-8
                    content = file_stream.getvalue().decode('utf-8')

            except UnicodeDecodeError:
                #падение при кодировке
                try:
                    content = file_stream.getvalue().decode('cp1251')
                except Exception as e:
                    flash(f"Не удалось определить кодировку текстового файла: {str(e)}", "error")
                    return render_template("documents/upload.html", project=project)
            except Exception as e:
                flash(f"Не удалось извлечь текст из файла: {str(e)}", "error")
                return render_template("documents/upload.html", project=project)

        #если все пустое
        if not content:
            flash("Пожалуйста, введите текст вручную или прикрепите документ (.docx, .txt)", "error")
            return render_template("documents/upload.html", project=project)

        #сохранение
        document = Document(
            name=name,
            content=content,
            project_id=project.id
        )
        db.session.add(document)
        db.session.commit()

        #аннотирование
        annotate_document(document.id)

        return redirect(url_for("documents.document_detail", document_id=document.id))

    return render_template("documents/upload.html", project=project)


@documents_bp.route("/<int:document_id>")
def document_detail(document_id):
    document = Document.query.get_or_404(document_id)
    highlighted_text = build_highlighted_text(document)
    return render_template(
        "documents/detail.html",
        document=document,
        highlighted_text=highlighted_text
    )


@documents_bp.route('/document/<int:document_id>/export/csv')
def export_to_csv(document_id):
    document = Document.query.get_or_404(document_id)
    response = export_csv(document, document_id)
    return response


@documents_bp.route('/document/<int:document_id>/export/excel')
def export_to_excel(document_id):
    document = Document.query.get_or_404(document_id)
    return export_excel(document, document_id)